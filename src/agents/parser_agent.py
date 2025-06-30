"""Parser Agent for extracting text and structure from documents."""

import os
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from crewai import Agent, Task, Crew
from crewai_tools import BaseTool
from pydantic import BaseModel, Field
import json
from pathlib import Path
from src.utils.file_manager import DocumentFileManager


class DocumentStructure(BaseModel):
    """Model for document structure data with semantic analysis."""
    
    paragraphs: List[Dict[str, Any]] = Field(description="List of paragraphs with text and formatting")
    tables: List[Dict[str, Any]] = Field(description="List of tables with content and formatting")
    metadata: Dict[str, Any] = Field(description="Document metadata like fonts, styles, etc.")
    semantic_structure: Dict[str, Any] = Field(description="Semantic document structure analysis")
    document_outline: List[Dict[str, Any]] = Field(description="Hierarchical document outline")
    
    
class DocumentParserTool(BaseTool):
    """Tool for parsing DOCX and TXT documents."""
    
    name: str = "document_parser"
    description: str = "Parse DOCX and TXT documents to extract text, structure, and formatting metadata"
    
    def _run(self, file_path: str, save_to_files: bool = True) -> str:
        """Parse a document and return structured data."""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return f"Error: File {file_path} does not exist"
            
            # Parse the document
            if file_path.suffix.lower() == '.docx':
                result_json = self._parse_docx(file_path)
            elif file_path.suffix.lower() == '.txt':
                result_json = self._parse_txt(file_path)
            else:
                return f"Error: Unsupported file format {file_path.suffix}"
            
            # Save to files if requested
            if save_to_files:
                try:
                    file_manager = DocumentFileManager()
                    session_id = file_manager.generate_session_id(str(file_path))
                    
                    # Parse the JSON result to save it
                    parsed_data = json.loads(result_json)
                    
                    # Save all files
                    saved_files = file_manager.save_parser_output(
                        session_id=session_id,
                        parsed_data=parsed_data,
                        original_filename=file_path.name
                    )
                    
                    # Add session info to the result
                    enhanced_result = {
                        'parsing_result': parsed_data,
                        'session_info': {
                            'session_id': session_id,
                            'saved_files': saved_files,
                            'status': 'parsed'
                        }
                    }
                    
                    return json.dumps(enhanced_result, indent=2, ensure_ascii=False)
                    
                except Exception as save_error:
                    # If saving fails, return original result with warning
                    return f"{result_json}\n\n[WARNING: Failed to save files: {str(save_error)}]"
            
            return result_json
                
        except Exception as e:
            return f"Error parsing document: {str(e)}"
    
    def _parse_docx(self, file_path: Path) -> str:
        """Parse a DOCX document."""
        doc = Document(file_path)
        
        # Extract paragraphs and tables in document order
        paragraphs = []
        tables = []
        element_index = 0
        
        # Iterate through document elements in order to preserve positioning
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph element
                # Find corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element == element:
                        para_text = para.text.strip()
                        if para_text:  # Skip empty paragraphs
                            para_data = {
                                "text": para_text,
                                "style": para.style.name if para.style else "Normal",
                                "alignment": self._get_alignment(para.alignment),
                                "runs": [],
                                "index": element_index  # Assign actual document position
                            }
                            
                            # Extract run-level formatting
                            for run in para.runs:
                                run_data = {
                                    "text": run.text,
                                    "bold": run.bold if run.bold is not None else False,
                                    "italic": run.italic if run.italic is not None else False,
                                    "underline": run.underline if run.underline is not None else False,
                                    "font_name": run.font.name if run.font.name else "Default",
                                    "font_size": str(run.font.size) if run.font.size else "Default"
                                }
                                
                                # Extract color if available
                                if run.font.color and run.font.color.rgb:
                                    color = run.font.color.rgb
                                    run_data["font_color"] = f"#{color.red:02x}{color.green:02x}{color.blue:02x}"
                                else:
                                    run_data["font_color"] = "Default"
                                    
                                para_data["runs"].append(run_data)
                            
                            paragraphs.append(para_data)
                            element_index += 1
                        break
                        
            elif element.tag.endswith('tbl'):  # Table element
                # Find corresponding table object
                for table in doc.tables:
                    if table._element == element:
                        table_data = {
                            "rows": [],
                            "index": element_index  # Assign actual document position
                        }
                        
                        for row in table.rows:
                            row_data = {
                                "cells": []
                            }
                            
                            for cell in row.cells:
                                cell_data = {
                                    "text": cell.text,
                                    "paragraphs": []
                                }
                                
                                # Extract cell paragraph formatting
                                for para in cell.paragraphs:
                                    para_data = {
                                        "text": para.text,
                                        "style": para.style.name if para.style else "Normal",
                                        "alignment": self._get_alignment(para.alignment)
                                    }
                                    cell_data["paragraphs"].append(para_data)
                                
                                row_data["cells"].append(cell_data)
                            
                            table_data["rows"].append(row_data)
                        
                        tables.append(table_data)
                        element_index += 1
                        break

        # Extract document metadata
        metadata = {
            "filename": file_path.name,
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables),
            "styles_used": list(set([para["style"] for para in paragraphs])),
            "fonts_used": list(set([run["font_name"] for para in paragraphs for run in para["runs"]])),
            "has_formatting": any(
                run["bold"] or run["italic"] or run["underline"] or run["font_color"] != "Default"
                for para in paragraphs for run in para["runs"]
            )
        }
        
        # Perform semantic analysis
        semantic_structure = self._analyze_semantic_structure(paragraphs)
        document_outline = self._build_document_outline(paragraphs, semantic_structure)
        
        # Add semantic analysis to metadata
        metadata.update({
            'semantic_content_types': semantic_structure['content_types'],
            'has_document_title': semantic_structure['document_title'] is not None,
            'hierarchy_levels': len([h for h in semantic_structure['hierarchy_structure'] if h['level'] is not None]),
            'total_sections': len([h for h in semantic_structure['hierarchy_structure'] if h['type'] == 'main_header']),
            'total_subsections': len([h for h in semantic_structure['hierarchy_structure'] if h['type'] == 'sub_header'])
        })
        
        # Create enhanced document structure
        doc_structure = DocumentStructure(
            paragraphs=paragraphs,
            tables=tables,
            metadata=metadata,
            semantic_structure=semantic_structure,
            document_outline=document_outline
        )
        
        return doc_structure.model_dump_json(indent=2)
    
    def _parse_txt(self, file_path: Path) -> str:
        """Parse a TXT document."""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Split into paragraphs (separated by double newlines)
        paragraphs = content.split('\n\n')
        
        paragraphs_data = []
        for para_text in paragraphs:
            if para_text.strip():
                para_data = {
                    "text": para_text.strip(),
                    "style": "Normal",
                    "alignment": "left",
                    "runs": [{
                        "text": para_text.strip(),
                        "bold": False,
                        "italic": False,
                        "underline": False,
                        "font_name": "Default",
                        "font_size": "Default",
                        "font_color": "Default"
                    }]
                }
                paragraphs_data.append(para_data)
        
        metadata = {
            "filename": file_path.name,
            "total_paragraphs": len(paragraphs_data),
            "total_tables": 0,
            "styles_used": ["Normal"],
            "fonts_used": ["Default"],
            "has_formatting": False
        }
        
        # Perform basic semantic analysis for TXT files
        semantic_structure = self._analyze_semantic_structure(paragraphs_data)
        document_outline = self._build_document_outline(paragraphs_data, semantic_structure)
        
        # Add semantic analysis to metadata
        metadata.update({
            'semantic_content_types': semantic_structure['content_types'],
            'has_document_title': semantic_structure['document_title'] is not None,
            'hierarchy_levels': 0,  # TXT files don't have style-based hierarchy
            'total_sections': 0,
            'total_subsections': 0
        })
        
        doc_structure = DocumentStructure(
            paragraphs=paragraphs_data,
            tables=[],
            metadata=metadata,
            semantic_structure=semantic_structure,
            document_outline=document_outline
        )
        
        return doc_structure.model_dump_json(indent=2)
    
    def _get_alignment(self, alignment) -> str:
        """Convert paragraph alignment to string."""
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return "center"
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return "right"
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return "justify"
        else:
            return "left"
    
    def _analyze_semantic_structure(self, paragraphs: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Analyze semantic structure of the document."""
        
        # Define style-to-semantic mapping
        style_semantic_map = {
            'CustomTitle': {'type': 'document_title', 'level': 0, 'translatable': True},
            'CustomH1': {'type': 'main_header', 'level': 1, 'translatable': True},
            'CustomH2': {'type': 'sub_header', 'level': 2, 'translatable': True},
            'CustomH3': {'type': 'sub_header', 'level': 3, 'translatable': True},
            'Heading 1': {'type': 'main_header', 'level': 1, 'translatable': True},
            'Heading 2': {'type': 'sub_header', 'level': 2, 'translatable': True},
            'Heading 3': {'type': 'sub_header', 'level': 3, 'translatable': True},
            'Normal': {'type': 'body_text', 'level': None, 'translatable': True},
            'FixedContent': {'type': 'special_content', 'level': None, 'translatable': True},
            'List Paragraph': {'type': 'list_item', 'level': None, 'translatable': True},
            'Quote': {'type': 'quote', 'level': None, 'translatable': True},
            'Caption': {'type': 'caption', 'level': None, 'translatable': True}
        }
        
        semantic_analysis = {
            'document_title': None,
            'main_sections': [],
            'content_types': {},
            'hierarchy_structure': []
        }
        
        # Analyze each paragraph for semantic meaning
        for i, para in enumerate(paragraphs):
            style = para.get('style', 'Normal')
            text = para.get('text', '').strip()
            
            if not text:  # Skip empty paragraphs
                continue
                
            # Get semantic information
            semantic_info = style_semantic_map.get(style, {
                'type': 'unknown', 'level': None, 'translatable': True
            })
            
            # Enhanced paragraph with semantic info
            para['semantic_type'] = semantic_info['type']
            para['hierarchy_level'] = semantic_info['level']
            para['is_translatable'] = semantic_info['translatable']
            para['paragraph_index'] = i
            
            # Track document title
            if semantic_info['type'] == 'document_title':
                semantic_analysis['document_title'] = {
                    'text': text,
                    'index': i,
                    'style': style
                }
            
            # Track content types
            content_type = semantic_info['type']
            if content_type not in semantic_analysis['content_types']:
                semantic_analysis['content_types'][content_type] = 0
            semantic_analysis['content_types'][content_type] += 1
            
            # Build hierarchy structure
            if semantic_info['level'] is not None:
                semantic_analysis['hierarchy_structure'].append({
                    'text': text,
                    'level': semantic_info['level'],
                    'index': i,
                    'type': semantic_info['type'],
                    'style': style
                })
        
        return semantic_analysis
    
    def _build_document_outline(self, paragraphs: List[Dict[str, Any]], semantic_structure: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Build hierarchical document outline."""
        
        outline = []
        current_section = None
        current_subsection = None
        
        for para in paragraphs:
            text = para.get('text', '').strip()
            if not text:
                continue
                
            semantic_type = para.get('semantic_type')
            level = para.get('hierarchy_level')
            
            if semantic_type == 'document_title':
                outline.append({
                    'type': 'title',
                    'text': text,
                    'level': 0,
                    'index': para.get('paragraph_index'),
                    'style': para.get('style'),
                    'children': []
                })
                
            elif semantic_type == 'main_header':
                current_section = {
                    'type': 'section',
                    'text': text,
                    'level': level,
                    'index': para.get('paragraph_index'),
                    'style': para.get('style'),
                    'children': [],
                    'content': []
                }
                outline.append(current_section)
                current_subsection = None
                
            elif semantic_type == 'sub_header':
                subsection = {
                    'type': 'subsection',
                    'text': text,
                    'level': level,
                    'index': para.get('paragraph_index'),
                    'style': para.get('style'),
                    'children': [],
                    'content': []
                }
                
                if current_section:
                    current_section['children'].append(subsection)
                else:
                    outline.append(subsection)
                current_subsection = subsection
                
            elif semantic_type in ['body_text', 'special_content', 'list_item']:
                content_item = {
                    'type': 'content',
                    'text': text[:100] + '...' if len(text) > 100 else text,  # Truncate for outline
                    'full_text': text,
                    'semantic_type': semantic_type,
                    'index': para.get('paragraph_index'),
                    'style': para.get('style')
                }
                
                # Add to appropriate parent
                if current_subsection:
                    current_subsection['content'].append(content_item)
                elif current_section:
                    current_section['content'].append(content_item)
                else:
                    # Orphaned content - add to outline directly
                    outline.append(content_item)
        
        return outline


def create_parser_agent(llm=None) -> Agent:
    """Create and configure the Parser Agent."""
    
    parser_tool = DocumentParserTool()
    
    # Create agent configuration
    agent_config = {
        "role": "Document Parser Specialist",
        "goal": "Extract text, structure, and formatting metadata from documents while preserving all original formatting information",
        "backstory": """You are an expert document parser with deep knowledge of document structures, 
        formatting, and metadata extraction. You specialize in analyzing DOCX and TXT files to extract 
        not just the text content, but also the complete formatting information including fonts, colors, 
        styles, alignments, and structural elements like tables and paragraphs. Your parsing ensures 
        that no formatting detail is lost, making it possible to reconstruct the document later.""",
        "tools": [parser_tool],
        "verbose": True,
        "allow_delegation": False
    }
    
    # Add LLM only if provided
    if llm is not None:
        agent_config["llm"] = llm
    
    agent = Agent(**agent_config)
    
    return agent


def create_parsing_task(agent: Agent, file_path: str) -> Task:
    """Create a parsing task for a specific document."""
    
    return Task(
        description=f"""Parse the document at '{file_path}' and extract:
        1. All text content with paragraph-level structure
        2. Complete formatting metadata including fonts, colors, styles, and alignments
        3. Table structures with cell-level formatting
        4. Document metadata and statistics
        5. A complete JSON structure that preserves all formatting information
        
        Ensure that the extracted data is comprehensive enough to reconstruct the document 
        with identical formatting later in the pipeline.""",
        expected_output="""A detailed JSON structure containing:
        - paragraphs: Array of paragraph objects with text and formatting
        - tables: Array of table objects with cell-level content and formatting
        - metadata: Document statistics and formatting summary
        The JSON should be properly formatted and include all necessary details for document reconstruction.""",
        agent=agent
    ) 