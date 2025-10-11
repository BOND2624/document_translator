"""
Output Agent - Generates DOCX documents from styled content
Handles document creation with proper formatting, fonts, colors, and structure
"""

import json
import os
from datetime import datetime
from typing import Dict, List, Any, Optional, Union
from pathlib import Path

from crewai import Agent, Task, Crew, Process
from crewai_tools import BaseTool
from pydantic import BaseModel, Field
from typing import ClassVar

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from ..utils.file_manager import DocumentFileManager


class DocumentOutputTool(BaseTool):
    """Tool for generating DOCX documents from styled content"""
    
    name: str = "document_output_tool" 
    description: str = "Generates DOCX documents from styled content with proper formatting and structure"
    
    def __init__(self):
        super().__init__()
    
    def _hex_to_rgb(self, hex_color: str) -> RGBColor:
        """Convert hex color to RGBColor object"""
        if hex_color.startswith('#'):
            hex_color = hex_color[1:]
        
        try:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16) 
            b = int(hex_color[4:6], 16)
            return RGBColor(r, g, b)
        except:
            # Default to black if color parsing fails
            return RGBColor(0, 0, 0)
    
    def _parse_spacing(self, spacing_str: str) -> Dict[str, float]:
        """Parse spacing string like '18pt before, 6pt after' into values"""
        spacing = {"before": 0, "after": 0}
        
        if not spacing_str:
            return spacing
        
        parts = spacing_str.lower().split(',')
        for part in parts:
            part = part.strip()
            if 'before' in part:
                value = part.replace('pt', '').replace('before', '').strip()
                try:
                    spacing["before"] = float(value)
                except:
                    spacing["before"] = 0
            elif 'after' in part:
                value = part.replace('pt', '').replace('after', '').strip()
                try:
                    spacing["after"] = float(value)
                except:
                    spacing["after"] = 0
        
        return spacing
    
    def _get_alignment(self, alignment_str: str) -> WD_PARAGRAPH_ALIGNMENT:
        """Convert alignment string to docx alignment enum"""
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        return alignment_map.get(alignment_str.lower(), WD_ALIGN_PARAGRAPH.LEFT)
    
    def _apply_run_formatting(self, run, run_data: Dict[str, Any], default_font: str = "Arial"):
        """Apply formatting to a text run"""
        run.text = run_data.get("text", "")
        self._apply_run_formatting_without_text(run, run_data, default_font)
    
    def _apply_run_formatting_without_text(self, run, run_data: Dict[str, Any], default_font: str = "Arial"):
        """Apply formatting to a text run without setting text"""
        # Font formatting
        run.font.name = run_data.get("font_name", default_font)
        run.font.bold = run_data.get("bold", False)
        run.font.italic = run_data.get("italic", False)
        run.font.underline = run_data.get("underline", False)
        
        # Font size
        font_size = run_data.get("font_size", "11pt")
        if font_size != "Default":
            try:
                size_pt = float(font_size.replace("pt", ""))
                run.font.size = Pt(size_pt)
            except:
                run.font.size = Pt(11)
        
        # Font color
        font_color = run_data.get("font_color", "#000000")
        if font_color != "Default":
            run.font.color.rgb = self._hex_to_rgb(font_color)
    
    def _add_paragraph_with_formatting(self, doc: Document, element: Dict[str, Any], 
                                     style_profile: Dict[str, Any]) -> None:
        """Add a paragraph with proper formatting"""
        
        # Get styled formatting or fall back to original
        formatting = element.get("styled_formatting", element.get("formatting", {}))
        
        # Create paragraph
        paragraph = doc.add_paragraph()
        
        # Set alignment
        alignment = formatting.get("alignment", "left")
        paragraph.alignment = self._get_alignment(alignment)
        
        
        # Set spacing
        spacing_str = formatting.get("spacing", "6pt after")
        spacing = self._parse_spacing(spacing_str)
        
        if spacing["before"] > 0:
            paragraph.paragraph_format.space_before = Pt(spacing["before"])
        if spacing["after"] > 0:
            paragraph.paragraph_format.space_after = Pt(spacing["after"])
        
        # Add text runs with formatting
        runs = formatting.get("runs", [])
        translated_text = element.get("translated_text", element.get("original_text", ""))
        
        if runs and len(runs) > 1:
            # Multiple runs - need to split translated text appropriately
            # For now, use the complete translated text with formatting from first run
            run = paragraph.add_run(translated_text)
            
            # Apply formatting from the most significant run (usually the first one with special formatting)
            primary_run = runs[0]
            for run_data in runs:
                if run_data.get("bold", False) or run_data.get("italic", False) or run_data.get("font_color", "Default") != "Default":
                    primary_run = run_data
                    break
            
            self._apply_run_formatting_without_text(run, primary_run,
                                                  style_profile.get("font_families", ["Arial"])[0])
        elif runs:
            # Single run - use translated text with original formatting
            run = paragraph.add_run(translated_text)
            self._apply_run_formatting_without_text(run, runs[0],
                                                  style_profile.get("font_families", ["Arial"])[0])
        else:
            # No runs, add text directly
            if translated_text:
                run = paragraph.add_run(translated_text)
                # Apply default formatting
                semantic_type = element.get("semantic_type", "body_text")
                font_family = style_profile.get("font_families", ["Arial"])[0]
                run.font.name = font_family
                
                # Apply semantic-based formatting
                if semantic_type == "document_title":
                    run.font.size = Pt(18)
                    run.font.bold = True
                    run.font.color.rgb = self._hex_to_rgb("#2c3e50")
                elif semantic_type in ["main_header", "sub_header"]:
                    run.font.size = Pt(14 if semantic_type == "main_header" else 12)
                    run.font.bold = True
                    run.font.color.rgb = self._hex_to_rgb("#34495e")
                else:
                    run.font.size = Pt(11)
                    run.font.color.rgb = self._hex_to_rgb("#1a1a1a")
    
    def _add_table_with_formatting(self, doc: Document, table_data: Dict[str, Any], 
                                   style_profile: Dict[str, Any]) -> None:
        """Add a table with proper formatting"""
        
        # Get the table structure from styled content
        original_table = table_data.get("original_table", {})
        rows_data = original_table.get("rows", [])
        
        if not rows_data:
            return
        
        # Determine table dimensions
        num_rows = len(rows_data)
        num_cols = len(rows_data[0].get("cells", [])) if rows_data else 0
        
        if num_cols == 0:
            return
        
        # Create table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Fill table data
        for row_idx, row_data in enumerate(rows_data):
            cells_data = row_data.get("cells", [])
            for col_idx, cell_data in enumerate(cells_data):
                if col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    # Extract text from cell data structure
                    cell_text = cell_data.get("text", "")
                    cell.text = str(cell_text)
                    
                    # Format cell paragraphs
                    for paragraph in cell.paragraphs:
                        # Format first row as header
                        if row_idx == 0:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                run.font.size = Pt(10)
                        else:
                            # Format data rows
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
        
        # Add some spacing after table
        doc.add_paragraph()
    
    def _create_document_structure(self, styled_elements: Dict[str, Any], 
                                 style_profile: Dict[str, Any]) -> Document:
        """Create the main document structure with all elements"""
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Add document title
        if "title" in styled_elements and styled_elements["title"]:
            self._add_paragraph_with_formatting(doc, styled_elements["title"], style_profile)
            # Add space after title
            doc.add_paragraph()
        
        # Add headers, content, and tables in order
        all_elements = []
        
        # Collect all elements with their indices for proper ordering
        if "headers" in styled_elements:
            for header_type in ["main_headers", "sub_headers"]:
                if header_type in styled_elements["headers"]:
                    for header in styled_elements["headers"][header_type]:
                        header["element_type"] = "paragraph"
                        all_elements.append(header)
        
        if "content" in styled_elements:
            for content_type in ["body_text", "special_content"]:
                if content_type in styled_elements["content"]:
                    for content_item in styled_elements["content"][content_type]:
                        content_item["element_type"] = "paragraph"
                        all_elements.append(content_item)
        
        # Add tables to the elements list
        if "tables" in styled_elements:
            for table_data in styled_elements["tables"]:
                table_data["element_type"] = "table"
                # Extract index from original_table if it exists
                if "original_table" in table_data and "index" in table_data["original_table"]:
                    table_data["index"] = table_data["original_table"]["index"]
                all_elements.append(table_data)
        
        # Sort by original index to maintain document order
        all_elements.sort(key=lambda x: x.get("index", 0))
        
        # Add elements to document in correct order
        for element in all_elements:
            if element.get("element_type") == "table":
                self._add_table_with_formatting(doc, element, style_profile)
            else:
                self._add_paragraph_with_formatting(doc, element, style_profile)
        
        return doc
    
    def _run(self, session_id: str, output_filename: Optional[str] = None) -> str:
        """Generate DOCX document from styled content"""
        
        try:
            file_manager = DocumentFileManager()
            
            # Load styled content from Style Agent
            styled_content_path = file_manager.get_session_file(session_id, "03_output_agent_input.json")
            if not os.path.exists(styled_content_path):
                return f"Error: Styled content not found for session {session_id}"
            
            with open(styled_content_path, 'r', encoding='utf-8') as f:
                styled_data = json.load(f)
            
            # Extract key information
            session_metadata = styled_data.get("session_metadata", {})
            generation_context = styled_data.get("generation_context", {})
            document_structure = styled_data.get("document_structure", {})
            style_specifications = styled_data.get("style_specifications", {})
            
            # Create the document
            doc = self._create_document_structure(document_structure, style_specifications)
            
            # Determine output filename
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                language = generation_context.get("language", "translated")
                output_filename = f"translated_document_{language}_{timestamp}.docx"
            
            # Save document
            session_dir = Path(file_manager.base_output_dir) / session_id
            output_path = session_dir / output_filename
            doc.save(str(output_path))
            
            # Create generation summary
            generation_summary = {
                "generation_metadata": {
                    "session_id": session_id,
                    "generated_at": datetime.now().isoformat(),
                    "output_file": output_filename,
                    "language": generation_context.get("language", "unknown"),
                    "style_profile": generation_context.get("style_profile", "unknown")
                },
                "document_stats": {
                    "total_elements": len(document_structure.get("headers", {}).get("main_headers", [])) + 
                                   len(document_structure.get("headers", {}).get("sub_headers", [])) + 
                                   len(document_structure.get("content", {}).get("body_text", [])) + 
                                   len(document_structure.get("content", {}).get("special_content", [])),
                    "has_title": document_structure.get("title") is not None,
                    "main_headers": len(document_structure.get("headers", {}).get("main_headers", [])),
                    "sub_headers": len(document_structure.get("headers", {}).get("sub_headers", [])),
                    "body_paragraphs": len(document_structure.get("content", {}).get("body_text", [])),
                    "special_content": len(document_structure.get("content", {}).get("special_content", [])),
                    "tables": len(document_structure.get("tables", []))
                },
                "formatting_applied": {
                    "fonts": style_specifications.get("fonts", []),
                    "colors": style_specifications.get("colors", {}),
                    "alignments": style_specifications.get("alignments", {}),
                    "spacing": style_specifications.get("spacing", {})
                }
            }
            
            # Save generation summary
            summary_path = file_manager.get_session_file(session_id, "04_generation_summary.json")
            with open(summary_path, 'w', encoding='utf-8') as f:
                json.dump(generation_summary, f, indent=2, ensure_ascii=False)
            
            # Create QA Agent input
            qa_input = {
                "session_metadata": {
                    **session_metadata,
                    "document_generation_completed_at": datetime.now().isoformat(),
                    "ready_for_qa_agent": True
                },
                "generated_document": {
                    "file_path": str(output_path),
                    "filename": output_filename,
                    "file_size_bytes": os.path.getsize(output_path),
                    "language": generation_context.get("language", "unknown")
                },
                "quality_check_requirements": {
                    "check_formatting_consistency": True,
                    "verify_translation_accuracy": True,
                    "validate_document_structure": True,
                    "assess_readability": True,
                    "check_language_compliance": True
                },
                "original_context": {
                    "source_language": generation_context.get("source_language", "unknown"),
                    "target_language": generation_context.get("language", "unknown"),
                    "document_type": generation_context.get("document_type", "unknown"),
                    "style_profile": generation_context.get("style_profile", "unknown")
                }
            }
            
            qa_input_path = file_manager.get_session_file(session_id, "04_qa_agent_input.json")
            with open(qa_input_path, 'w', encoding='utf-8') as f:
                json.dump(qa_input, f, indent=2, ensure_ascii=False)
            
            # Update session manifest
            file_manager.update_session_manifest(session_id, {
                "output_agent_completed": True,
                "generated_document": output_filename,
                "generation_summary_file": "04_generation_summary.json",
                "qa_agent_input_file": "04_qa_agent_input.json",
                "document_file_size": os.path.getsize(output_path),
                "generation_stage": "completed"
            })
            
            file_size_mb = os.path.getsize(output_path) / (1024 * 1024)
            total_elements = generation_summary["document_stats"]["total_elements"]
            
            return f"Document generation completed successfully. Created '{output_filename}' ({file_size_mb:.2f} MB) with {total_elements} elements. Files saved: generation_summary, qa_agent_input. Ready for QA validation."
            
        except Exception as e:
            return f"Error in document generation: {str(e)}"


def create_output_agent(llm=None) -> Agent:
    """Create the Output Agent with proper configuration"""
    
    output_tool = DocumentOutputTool()
    
    # Create agent configuration
    agent_config = {
        "role": 'Document Generation Specialist',
        "goal": 'Generate professional DOCX documents from styled content while preserving formatting, structure, and language-specific requirements',
        "backstory": """You are an expert in document generation and formatting. 
        You specialize in creating professional Word documents that maintain visual consistency, 
        proper typography, and language-appropriate formatting for LTR languages. Your expertise includes 
        handling complex document structures, multi-language text rendering, 
        and ensuring that translated content appears professional and polished.""",
        "tools": [output_tool],
        "verbose": True,
        "allow_delegation": False
    }
    
    # Add LLM only if provided
    if llm is not None:
        agent_config["llm"] = llm
    
    return Agent(**agent_config)


def create_output_task(agent: Agent, session_id: str, output_filename: Optional[str] = None) -> Task:
    """Create the document generation task"""
    
    return Task(
        description=f"""Generate a professional DOCX document from the styled content in session {session_id}.

        Your responsibilities:
        1. Load the styled content from the Style Agent output
        2. Create a properly formatted Word document with all styling applied
        3. Preserve document structure and semantic hierarchy  
        4. Apply LTR language-specific formatting requirements
        5. Ensure professional appearance and readability
        6. Generate document statistics and quality metrics
        7. Prepare input for the QA Agent
        
        Use the document_output_tool to generate the document{f' with filename: {output_filename}' if output_filename else ''}.
        """,
        
        expected_output="""Complete document generation with:
        1. Professional DOCX document with proper formatting applied
        2. Generation summary with document statistics and metadata
        3. QA Agent input file with quality check requirements
        4. Updated session manifest with generation status
        
        The generated document should be ready for quality validation.""",
        
        agent=agent
    )


def run_output_agent(session_id: str, output_filename: Optional[str] = None) -> str:
    """Run the Output Agent workflow"""
    
    # Create LLM and agent with it
    from ..config import Config
    llm = Config.create_llm()
    output_agent = create_output_agent(llm=llm)
    output_task = create_output_task(output_agent, session_id, output_filename)
    
    # Create and run crew
    crew = Crew(
        agents=[output_agent],
        tasks=[output_task],
        process=Process.sequential,
        verbose=True
    )
    
    # Execute the workflow
    result = crew.kickoff()
    
    return str(result)


if __name__ == "__main__":
    # Example usage
    session_id = "d212cbd859d2_20250621_125824"  # Replace with actual session ID
    result = run_output_agent(session_id)
    print("Output Agent Result:", result) 