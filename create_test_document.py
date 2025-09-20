"""
Test Document Generator for Document Translation System
Creates a comprehensive DOCX document with various elements for testing
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import os

def create_test_document():
    """Create a comprehensive test document with various elements"""
    
    # Create new document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # ========================================
    # DOCUMENT TITLE
    # ========================================
    title = doc.add_heading('Global Technology Solutions Annual Report 2024', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Strategic Innovation and Market Leadership Analysis')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(70, 70, 70)
    subtitle_run.italic = True
    
    # Add spacing
    doc.add_paragraph()
    
    # ========================================
    # EXECUTIVE SUMMARY
    # ========================================
    doc.add_heading('Executive Summary', level=1)
    
    exec_summary = doc.add_paragraph(
        "Global Technology Solutions achieved exceptional performance in 2024, establishing "
        "new benchmarks in innovation and market expansion. Our comprehensive strategy focused "
        "on artificial intelligence, cloud computing, and sustainable technology solutions has "
        "delivered remarkable results across all business segments."
    )
    exec_summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add key highlights with formatting
    highlights = doc.add_paragraph()
    highlights.add_run("Key Achievements: ").bold = True
    highlights.add_run(
        "Launched revolutionary AI platform, secured 15 major enterprise clients, "
        "expanded to Asian markets, and achieved industry-leading customer retention rate of 97%."
    )
    
    # ========================================
    # FINANCIAL PERFORMANCE
    # ========================================
    doc.add_heading('Financial Performance', level=1)
    
    doc.add_heading('Revenue Analysis', level=2)
    
    revenue_para = doc.add_paragraph(
        "Our technology solutions division generated exceptional financial returns in 2024. "
        "Total revenue reached $185.2 million, representing a 32% increase from 2023's "
        "$140.3 million. This growth was primarily driven by our AI platform subscriptions, "
        "cloud infrastructure services, and strategic partnerships with Fortune 500 companies."
    )
    revenue_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add financial table
    doc.add_heading('Quarterly Financial Summary', level=3)
    
    # Create financial table
    financial_table = doc.add_table(rows=5, cols=4)
    financial_table.style = 'Table Grid'
    financial_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = financial_table.rows[0].cells
    header_cells[0].text = 'Quarter'
    header_cells[1].text = 'Revenue (millions)'
    header_cells[2].text = 'Profit Margin (%)'
    header_cells[3].text = 'Growth Rate (%)'
    
    # Format header row
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Set header background (simplified - actual background setting is complex in python-docx)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    financial_data = [
        ['Q1 2024', '$42.1', '23.5%', '28.2%'],
        ['Q2 2024', '$45.8', '25.1%', '31.4%'],
        ['Q3 2024', '$48.3', '26.8%', '34.7%'],
        ['Q4 2024', '$49.0', '28.2%', '36.1%']
    ]
    
    for i, row_data in enumerate(financial_data, 1):
        cells = financial_table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========================================
    # MARKET ANALYSIS
    # ========================================
    doc.add_heading('Market Analysis and Competitive Landscape', level=1)
    
    doc.add_heading('Industry Overview', level=2)
    
    market_para = doc.add_paragraph(
        "The enterprise technology market demonstrates robust growth prospects, with AI and "
        "cloud services leading digital transformation initiatives. Our market position strengthens "
        "through strategic focus on emerging technologies including machine learning, edge computing, "
        "and cybersecurity solutions for enterprise clients."
    )
    
    # Add bullet points
    doc.add_heading('Key Market Trends', level=3)
    
    trends = [
        "Enterprise AI adoption accelerating at unprecedented pace",
        "Cloud-first strategies becoming standard across industries", 
        "Cybersecurity investments increasing due to rising threats",
        "Edge computing enabling real-time data processing",
        "Integration of IoT devices driving new service opportunities"
    ]
    
    for trend in trends:
        bullet_para = doc.add_paragraph(trend, style='List Bullet')
        bullet_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # ========================================
    # PRODUCT DEVELOPMENT
    # ========================================
    doc.add_heading('Product Development and Innovation', level=1)
    
    doc.add_heading('New Product Launches', level=2)
    
    product_para = doc.add_paragraph(
        "Our product development team delivered breakthrough innovations in 2024, including "
        "the next-generation AI analytics platform and advanced cybersecurity suite. "
        "Research and development investments totaled $22.8 million, representing 12.3% of "
        "total revenue, resulting in 8 new patents and 5 major product launches."
    )
    product_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Product portfolio table
    doc.add_heading('Product Portfolio Performance', level=3)
    
    product_table = doc.add_table(rows=6, cols=5)
    product_table.style = 'Table Grid'
    
    # Product table headers
    product_headers = ['Product Category', 'Units Sold', 'Revenue Share', 'Customer Rating', 'Market Position']
    header_row = product_table.rows[0].cells
    for i, header in enumerate(product_headers):
        header_row[i].text = header
        header_row[i].paragraphs[0].runs[0].font.bold = True
        header_row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Product data
    product_data = [
        ['AI Analytics Platform', '12,500', '42%', '4.9/5.0', '#1 in Segment'],
        ['Cloud Infrastructure', '8,200', '31%', '4.7/5.0', '#2 in Market'],
        ['Cybersecurity Suite', '5,400', '18%', '4.8/5.0', '#1 in Security'],
        ['Enterprise Consulting', '850', '6%', '4.9/5.0', '#1 in Advisory'],
        ['Training & Certification', '2,100', '3%', '4.6/5.0', '#3 in Education']
    ]
    
    for i, row_data in enumerate(product_data, 1):
        cells = product_table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========================================
    # OPERATIONAL EXCELLENCE
    # ========================================
    doc.add_heading('Operational Excellence and Process Optimization', level=1)
    
    doc.add_heading('Efficiency Improvements', level=2)
    
    operations_para = doc.add_paragraph(
        "Our commitment to operational excellence has resulted in significant improvements "
        "across all business functions. We implemented advanced automation systems, "
        "streamlined workflows, and adopted best practices that increased productivity "
        "by 35% while reducing operational costs by 18%."
    )
    
    # Add formatted text with special styling
    special_para = doc.add_paragraph()
    special_run = special_para.add_run("IMPORTANT NOTE: ")
    special_run.font.bold = True
    special_run.font.color.rgb = RGBColor(200, 0, 0)
    special_run.font.size = Pt(12)
    
    special_para.add_run(
        "All operational improvements were achieved while maintaining our commitment "
        "to employee welfare and environmental sustainability."
    )
    
    # ========================================
    # SUSTAINABILITY INITIATIVES
    # ========================================
    doc.add_heading('Environmental Sustainability and Social Responsibility', level=1)
    
    doc.add_heading('Carbon Neutrality Achievement', level=2)
    
    sustainability_para = doc.add_paragraph(
        "We are proud to announce that our organization achieved carbon neutrality six months "
        "ahead of our scheduled target. This milestone represents a significant step in our "
        "journey toward becoming a fully sustainable enterprise by 2026."
    )
    
    # Sustainability metrics table
    doc.add_heading('Environmental Impact Metrics', level=3)
    
    sustainability_table = doc.add_table(rows=5, cols=3)
    sustainability_table.style = 'Table Grid'
    
    # Sustainability headers
    sust_headers = ['Environmental Metric', '2023 Baseline', '2024 Achievement']
    sust_header_row = sustainability_table.rows[0].cells
    for i, header in enumerate(sust_headers):
        sust_header_row[i].text = header
        sust_header_row[i].paragraphs[0].runs[0].font.bold = True
    
    # Sustainability data
    sust_data = [
        ['Carbon Emissions (tons CO2)', '2,450', '0 (Net Zero)'],
        ['Energy Consumption (MWh)', '3,200', '2,100 (34% reduction)'],
        ['Renewable Energy Usage', '45%', '100%'],
        ['Waste Recycling Rate', '72%', '95%']
    ]
    
    for i, row_data in enumerate(sust_data, 1):
        cells = sustainability_table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
    
    # ========================================
    # FUTURE OUTLOOK
    # ========================================
    doc.add_heading('Strategic Outlook and Future Plans', level=1)
    
    doc.add_heading('2025 Strategic Priorities', level=2)
    
    outlook_para = doc.add_paragraph(
        "Looking ahead to 2025, we have identified key strategic priorities that will "
        "drive continued growth and innovation. Our focus remains on expanding our "
        "global presence, investing in cutting-edge technologies, and strengthening "
        "our commitment to sustainable business practices."
    )
    
    # Strategic priorities list
    doc.add_heading('Key Strategic Initiatives', level=3)
    
    priorities = [
        "Expand operations to 10 new international markets",
        "Launch next-generation AI-powered product suite",
        "Achieve 40% increase in renewable energy infrastructure",
        "Establish strategic partnerships with industry leaders",
        "Implement comprehensive digital transformation program"
    ]
    
    for i, priority in enumerate(priorities, 1):
        priority_para = doc.add_paragraph(f"{i}. {priority}")
        priority_para.style = 'List Number'
    
    # ========================================
    # CONCLUSION
    # ========================================
    doc.add_heading('Conclusion', level=1)
    
    conclusion_para = doc.add_paragraph(
        "The achievements of 2024 have solidified our position as an industry leader "
        "committed to innovation, sustainability, and excellence. As we move forward "
        "into 2025, we remain focused on delivering exceptional value to our stakeholders "
        "while maintaining our responsibility to the environment and society."
    )
    conclusion_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add closing statement
    closing_para = doc.add_paragraph()
    closing_run = closing_para.add_run("Thank you for your continued trust and partnership.")
    closing_run.font.bold = True
    closing_run.font.size = Pt(14)
    closing_run.font.color.rgb = RGBColor(0, 70, 140)
    closing_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add document footer information
    doc.add_paragraph()
    footer_para = doc.add_paragraph("Document prepared by: Strategic Planning Department")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(100, 100, 100)
    
    date_para = doc.add_paragraph("Last updated: December 2024")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.runs[0]
    date_run.font.size = Pt(10)
    date_run.font.italic = True
    date_run.font.color.rgb = RGBColor(100, 100, 100)
    
    return doc

def main():
    """Generate the test document and save it"""
    print("🔧 Creating comprehensive test document...")
    
    # Create the document
    doc = create_test_document()
    
    # Save the document
    filename = "test_document_comprehensive.docx"
    doc.save(filename)
    
    print(f"✅ Test document created: {filename}")
    print(f"📄 File size: {os.path.getsize(filename) / 1024:.1f} KB")
    
    # Print document statistics
    print("\n📊 Document Statistics:")
    print(f"   • Paragraphs: {len(doc.paragraphs)}")
    print(f"   • Tables: {len(doc.tables)}")
    print(f"   • Headers: Multiple levels (H1, H2, H3)")
    print(f"   • Content types: Title, headers, body text, tables, lists, formatted text")
    print(f"   • Special formatting: Bold, italic, colors, alignment, bullet points")
    
    print(f"\n🎯 Ready for testing with the translation system!")
    print(f"   Use this file to test: {filename}")

if __name__ == "__main__":
    main()
